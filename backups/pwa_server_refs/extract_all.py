"""
PWA Server Reference Data Extraction Script
Extracts all available reference data from PWA server folders.
"""
import json
import os
import sys
from pathlib import Path

import openpyxl
import pdfplumber
import xlrd
import docx

OUTPUT_DIR = Path(r"C:\Users\FacundoS-PC\dev\BarackMercosul\backups\pwa_server_refs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SERVER_BASE = r"\\SERVER\compartido\BARACK\CALIDAD\DOCUMENTACION SGC\PPAP CLIENTES\PWA"
PLANAS_BASE = os.path.join(SERVER_BASE, "1- TOYOTA_TELAS_ PLANAS_581D", "APQP")
TERMO_BASE = os.path.join(SERVER_BASE, "2-TOYOTA_TELAS_TERMOFORMADAS_582D", "APQP")
PPAP_BASE = os.path.join(SERVER_BASE, "PPAP - PWA PROYECTO (TELAS HILUX)")

summary = {
    "extraction_date": "2026-04-06",
    "files_extracted": [],
    "part_numbers": [],
    "errors": []
}


def safe_val(val):
    """Convert cell value to JSON-serializable form."""
    if val is None:
        return None
    if isinstance(val, (int, float, bool)):
        return val
    return str(val)


def extract_xlsx(filepath, output_name):
    """Extract all sheets from an xlsx file to JSON."""
    print(f"  Extracting XLSX: {os.path.basename(filepath)}")
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        result = {}
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=False):
                row_data = [safe_val(cell.value) for cell in row]
                rows.append(row_data)
            result[sheet_name] = rows

        out_path = OUTPUT_DIR / f"{output_name}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        summary["files_extracted"].append({
            "source": filepath,
            "output": str(out_path),
            "type": "xlsx",
            "sheets": wb.sheetnames,
            "status": "OK"
        })
        print(f"    -> Saved to {out_path}")
        return result
    except Exception as e:
        err = f"Error extracting {filepath}: {e}"
        print(f"    ERROR: {err}")
        summary["errors"].append(err)
        return None


def extract_xls(filepath, output_name):
    """Extract all sheets from an xls file to JSON."""
    print(f"  Extracting XLS: {os.path.basename(filepath)}")
    try:
        wb = xlrd.open_workbook(filepath)
        result = {}
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            rows = []
            for row_idx in range(ws.nrows):
                row_data = []
                for col_idx in range(ws.ncols):
                    cell = ws.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_DATE:
                        try:
                            date_tuple = xlrd.xldate_as_tuple(cell.value, wb.datemode)
                            row_data.append(f"{date_tuple[0]:04d}-{date_tuple[1]:02d}-{date_tuple[2]:02d}")
                        except:
                            row_data.append(safe_val(cell.value))
                    else:
                        row_data.append(safe_val(cell.value))
                rows.append(row_data)
            result[sheet_name] = rows

        out_path = OUTPUT_DIR / f"{output_name}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        summary["files_extracted"].append({
            "source": filepath,
            "output": str(out_path),
            "type": "xls",
            "sheets": wb.sheet_names(),
            "status": "OK"
        })
        print(f"    -> Saved to {out_path}")
        return result
    except Exception as e:
        err = f"Error extracting {filepath}: {e}"
        print(f"    ERROR: {err}")
        summary["errors"].append(err)
        return None


def extract_pdf(filepath, output_name):
    """Extract text and tables from a PDF file."""
    print(f"  Extracting PDF: {os.path.basename(filepath)}")
    try:
        result = {"pages": [], "tables": []}
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                result["pages"].append({
                    "page_num": i + 1,
                    "text": text
                })
                tables = page.extract_tables()
                if tables:
                    for t_idx, table in enumerate(tables):
                        clean_table = []
                        for row in table:
                            clean_table.append([safe_val(c) for c in row])
                        result["tables"].append({
                            "page": i + 1,
                            "table_index": t_idx,
                            "data": clean_table
                        })

        out_path = OUTPUT_DIR / f"{output_name}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        summary["files_extracted"].append({
            "source": filepath,
            "output": str(out_path),
            "type": "pdf",
            "pages": len(result["pages"]),
            "tables_found": len(result["tables"]),
            "status": "OK"
        })
        print(f"    -> Saved to {out_path} ({len(result['pages'])} pages, {len(result['tables'])} tables)")
        return result
    except Exception as e:
        err = f"Error extracting {filepath}: {e}"
        print(f"    ERROR: {err}")
        summary["errors"].append(err)
        return None


def extract_docx(filepath, output_name):
    """Extract text and tables from a DOCX file."""
    print(f"  Extracting DOCX: {os.path.basename(filepath)}")
    try:
        doc = docx.Document(filepath)
        result = {"paragraphs": [], "tables": []}

        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append({
                    "text": para.text,
                    "style": para.style.name if para.style else None
                })

        for t_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append({
                "table_index": t_idx,
                "data": table_data
            })

        out_path = OUTPUT_DIR / f"{output_name}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        summary["files_extracted"].append({
            "source": filepath,
            "output": str(out_path),
            "type": "docx",
            "paragraphs": len(result["paragraphs"]),
            "tables_found": len(result["tables"]),
            "status": "OK"
        })
        print(f"    -> Saved to {out_path}")
        return result
    except Exception as e:
        err = f"Error extracting {filepath}: {e}"
        print(f"    ERROR: {err}")
        summary["errors"].append(err)
        return None


def main():
    print("=" * 70)
    print("PWA Server Reference Data Extraction")
    print("=" * 70)

    # =========================================================================
    # 1. PLANAS 581D
    # =========================================================================
    print("\n--- PLANAS 581D ---")

    # 1a. Materials list
    print("\n[1] Materials list:")
    extract_xlsx(
        os.path.join(PLANAS_BASE, "7-Lista de materiales", "Listado de materiales.xlsx"),
        "planas_581d_materials_list"
    )

    # 1b. Flujograma
    print("\n[2] Flujograma:")
    extract_pdf(
        os.path.join(PLANAS_BASE, "9-Flujograma del proceso preliminar", "FLUJOGRAMA_150_TELAS_581D_REV.pdf"),
        "planas_581d_flujograma"
    )

    # 1c. Embalaje spec
    print("\n[3] Embalaje spec:")
    extract_pdf(
        os.path.join(PLANAS_BASE, "19-Normas y Especificaciones de embalaje", "FICHA DE EMBALAJE HILUX 581D REV.A.pdf"),
        "planas_581d_embalaje"
    )

    # 1d. Ficha tecnica files (DOCX)
    print("\n[4] Ficha tecnica files:")
    ft_dir = os.path.join(PLANAS_BASE, "8- Ficha tecnica")
    for fname in os.listdir(ft_dir):
        fpath = os.path.join(ft_dir, fname)
        if fname.lower().endswith(".docx") and not fname.startswith("~"):
            safe_name = fname.replace(" ", "_").replace(".", "_").rsplit("_", 1)[0]
            extract_docx(fpath, f"planas_581d_ft_{safe_name}")

    # 1e. Material weight data
    print("\n[5] Material weight data:")
    extract_xlsx(
        os.path.join(PLANAS_BASE, "14-Especificaciones de Materiales", "Peso de muestras 100 x 100 mm.xlsx"),
        "planas_581d_material_weights"
    )

    # 1f. s-felt PDF in material specs
    print("\n[6] S-felt material spec:")
    sfelt_path = os.path.join(PLANAS_BASE, "14-Especificaciones de Materiales", "s-felt-300.pdf")
    if os.path.exists(sfelt_path):
        extract_pdf(sfelt_path, "planas_581d_sfelt_300")

    # 1g. Materials list PDF version
    print("\n[7] Materials list (PDF version):")
    mat_pdf_path = os.path.join(PLANAS_BASE, "7-Lista de materiales", "Listado de materiales.pdf")
    if os.path.exists(mat_pdf_path):
        extract_pdf(mat_pdf_path, "planas_581d_materials_list_pdf")

    # =========================================================================
    # 2. TERMOFORMADAS 582D
    # =========================================================================
    print("\n--- TERMOFORMADAS 582D ---")

    # 2a. BOM
    print("\n[8] BOM:")
    extract_xlsx(
        os.path.join(TERMO_BASE, "7-Lista de materiales", "TOYOTA_TELAS_TERMOFORMADAS_582D_BOM Barack_Preliminar .xlsx"),
        "termo_582d_bom"
    )

    # 2b. Flujograma
    print("\n[9] Flujograma:")
    extract_pdf(
        os.path.join(TERMO_BASE, "9-Flujograma del proceso preliminar", "FLUJOGRAMA_156_ASIENTO_RESPALDO_REV.pdf"),
        "termo_582d_flujograma"
    )

    # 2c. Ficha tecnica PDFs
    print("\n[10] Ficha tecnica PDFs:")
    extract_pdf(
        os.path.join(TERMO_BASE, "8- Ficha tecnica", "FT153 Refuerzo Tricapa .pdf"),
        "termo_582d_ft153_refuerzo_tricapa"
    )
    extract_pdf(
        os.path.join(TERMO_BASE, "8- Ficha tecnica", "FT154 Refuerzo 1500 gm2 .pdf"),
        "termo_582d_ft154_refuerzo_1500"
    )

    # =========================================================================
    # 3. PPAP PROJECT
    # =========================================================================
    print("\n--- PPAP PROJECT ---")

    # 3a. Process Flow (folder 5) - check for files
    pf_dir = os.path.join(PPAP_BASE, "5- Process Flow")
    pf_files = os.listdir(pf_dir) if os.path.exists(pf_dir) else []
    print(f"\n[11] Process Flow folder: {pf_files if pf_files else 'EMPTY'}")
    summary["files_extracted"].append({
        "source": pf_dir,
        "type": "folder_check",
        "files_found": pf_files,
        "status": "EMPTY" if not pf_files else "HAS_FILES"
    })

    # 3b. PFMEA (folder 6) - check for files
    pfmea_dir = os.path.join(PPAP_BASE, "6- PFMEA")
    pfmea_files = os.listdir(pfmea_dir) if os.path.exists(pfmea_dir) else []
    print(f"\n[12] PFMEA folder: {pfmea_files if pfmea_files else 'EMPTY'}")
    summary["files_extracted"].append({
        "source": pfmea_dir,
        "type": "folder_check",
        "files_found": pfmea_files,
        "status": "EMPTY" if not pfmea_files else "HAS_FILES"
    })

    # 3c. Control Plan
    print("\n[13] Control Plan:")
    extract_xls(
        os.path.join(PPAP_BASE, "7- Control plan", "PC APOYABRAZOS P21 rev H 08-10-2024.xls"),
        "ppap_control_plan_apoyabrazos"
    )

    # =========================================================================
    # 4. PART NUMBERS FROM IMAGES FOLDER
    # =========================================================================
    print("\n--- PART NUMBERS FROM IMAGES ---")
    images_dir = os.path.join(SERVER_BASE, "1- TOYOTA_TELAS_ PLANAS_581D", "Imagenes pieza")
    part_numbers = []
    if os.path.exists(images_dir):
        for entry in sorted(os.listdir(images_dir)):
            full_path = os.path.join(images_dir, entry)
            if os.path.isdir(full_path):
                # List files inside each part number folder
                files_inside = os.listdir(full_path)
                pn_entry = {
                    "folder_name": entry,
                    "part_number_formatted": f"21-{entry[2:]}" if entry.startswith("21") and len(entry) == 6 else entry,
                    "files": files_inside
                }
                part_numbers.append(pn_entry)
                print(f"  {entry} -> {pn_entry['part_number_formatted']} ({len(files_inside)} files)")

    summary["part_numbers"] = part_numbers

    # Save part numbers separately
    pn_path = OUTPUT_DIR / "planas_581d_part_numbers.json"
    with open(pn_path, "w", encoding="utf-8") as f:
        json.dump(part_numbers, f, ensure_ascii=False, indent=2)
    print(f"  -> Saved part numbers to {pn_path}")

    # =========================================================================
    # 5. ALSO CHECK ADDITIONAL FOLDERS IN APQP
    # =========================================================================
    print("\n--- ADDITIONAL CHECKS ---")

    # Check Planas Plan de Control (folder 12)
    pc_planas_dir = os.path.join(PLANAS_BASE, "12-Plan de Control")
    if os.path.exists(pc_planas_dir):
        pc_files = os.listdir(pc_planas_dir)
        print(f"  Planas Plan de Control (12): {pc_files}")
        for fname in pc_files:
            fpath = os.path.join(pc_planas_dir, fname)
            if fname.lower().endswith(".xlsx") and not fname.startswith("~"):
                extract_xlsx(fpath, f"planas_581d_plan_control_{fname.replace(' ', '_').replace('.xlsx', '')}")
            elif fname.lower().endswith(".xls") and not fname.startswith("~"):
                extract_xls(fpath, f"planas_581d_plan_control_{fname.replace(' ', '_').replace('.xls', '')}")

    # Check Planas FMEA (folder 22)
    fmea_planas_dir = os.path.join(PLANAS_BASE, "22- FMEA de proceso")
    if os.path.exists(fmea_planas_dir):
        fmea_files = os.listdir(fmea_planas_dir)
        print(f"  Planas FMEA de proceso (22): {fmea_files}")
        for fname in fmea_files:
            fpath = os.path.join(fmea_planas_dir, fname)
            if fname.lower().endswith(".xlsx") and not fname.startswith("~"):
                extract_xlsx(fpath, f"planas_581d_fmea_{fname.replace(' ', '_').replace('.xlsx', '')}")
            elif fname.lower().endswith(".xls") and not fname.startswith("~"):
                extract_xls(fpath, f"planas_581d_fmea_{fname.replace(' ', '_').replace('.xls', '')}")
            elif fname.lower().endswith(".pdf"):
                extract_pdf(fpath, f"planas_581d_fmea_{fname.replace(' ', '_').replace('.pdf', '')}")

    # Check Planas Flujograma de proceso (folder 20)
    flujo_planas_dir = os.path.join(PLANAS_BASE, "20- Flujograma de proceso")
    if os.path.exists(flujo_planas_dir):
        flujo_files = os.listdir(flujo_planas_dir)
        print(f"  Planas Flujograma de proceso (20): {flujo_files}")
        for fname in flujo_files:
            fpath = os.path.join(flujo_planas_dir, fname)
            if fname.lower().endswith(".pdf"):
                extract_pdf(fpath, f"planas_581d_flujo20_{fname.replace(' ', '_').replace('.pdf', '')}")

    # Check Termoformadas Plan de Control (folder 12)
    pc_termo_dir = os.path.join(TERMO_BASE, "12-Plan de Control")
    if os.path.exists(pc_termo_dir):
        pc_files = os.listdir(pc_termo_dir)
        print(f"  Termoformadas Plan de Control (12): {pc_files}")
        for fname in pc_files:
            fpath = os.path.join(pc_termo_dir, fname)
            if fname.lower().endswith(".xlsx") and not fname.startswith("~"):
                extract_xlsx(fpath, f"termo_582d_plan_control_{fname.replace(' ', '_').replace('.xlsx', '')}")
            elif fname.lower().endswith(".xls") and not fname.startswith("~"):
                extract_xls(fpath, f"termo_582d_plan_control_{fname.replace(' ', '_').replace('.xls', '')}")

    # Check Termoformadas FMEA (folder 22)
    fmea_termo_dir = os.path.join(TERMO_BASE, "22- FMEA de proceso")
    if os.path.exists(fmea_termo_dir):
        fmea_files = os.listdir(fmea_termo_dir)
        print(f"  Termoformadas FMEA de proceso (22): {fmea_files}")
        for fname in fmea_files:
            fpath = os.path.join(fmea_termo_dir, fname)
            if fname.lower().endswith(".xlsx") and not fname.startswith("~"):
                extract_xlsx(fpath, f"termo_582d_fmea_{fname.replace(' ', '_').replace('.xlsx', '')}")
            elif fname.lower().endswith(".xls") and not fname.startswith("~"):
                extract_xls(fpath, f"termo_582d_fmea_{fname.replace(' ', '_').replace('.xls', '')}")

    # Check Termoformadas Flujograma de proceso (folder 20)
    flujo_termo_dir = os.path.join(TERMO_BASE, "20- Flujograma de proceso")
    if os.path.exists(flujo_termo_dir):
        flujo_files = os.listdir(flujo_termo_dir)
        print(f"  Termoformadas Flujograma de proceso (20): {flujo_files}")
        for fname in flujo_files:
            fpath = os.path.join(flujo_termo_dir, fname)
            if fname.lower().endswith(".pdf"):
                extract_pdf(fpath, f"termo_582d_flujo20_{fname.replace(' ', '_').replace('.pdf', '')}")

    # =========================================================================
    # SAVE SUMMARY
    # =========================================================================
    print("\n" + "=" * 70)
    print("SAVING SUMMARY")
    print("=" * 70)

    summary_path = OUTPUT_DIR / "extraction_summary.json"
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"\nSummary saved to: {summary_path}")
    print(f"Total files extracted: {len(summary['files_extracted'])}")
    print(f"Total errors: {len(summary['errors'])}")
    print(f"Part numbers found: {len(summary['part_numbers'])}")

    if summary["errors"]:
        print("\nERRORS:")
        for err in summary["errors"]:
            print(f"  - {err}")

    print("\nDone!")


if __name__ == "__main__":
    main()
